from docx import Document # type: ignore
import pandas as pd
import re
import logging

logging.basicConfig(
    level=logging.INFO,  # Mínimo nível de log que será exibido
    format='%(asctime)s - %(levelname)s - %(message)s'  # Formato das mensagens
)

import rules_and_definitions

def validate_data(data, pattern):
    return bool(re.match(pattern, data, re.IGNORECASE))


def classify_paragraph(text, pos_no_doc):
    match text:
        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_capitulo):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'capitulo'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_secao):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'secao'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_subsecao):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'subsecao'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_titulo_capitulo_secao_ou_subsecao):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'titulo-capitulo-secao-ou-subsecao'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_titulo):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'titulo'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_livro):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'livro'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_anexo):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'anexo'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_titulo_tabela):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'titulo-tabela'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_artigo):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'artigo'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_alinea):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'alinea'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_inciso):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'inciso'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_paragrafo):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'paragrafo'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return


        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_item):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'item'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_paragrafo_em_branco):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'paragrafo-em-branco'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_paragrafo_vazio):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'paragrafo-vazio'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_nota_explicativa):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'nota-explicativa'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_nota_de_rodape):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'nota-de-rodape'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_cfop):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'cfop'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_item_nota_de_rodape):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'item-nota-de-rodape'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_tabela):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'tabela'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_texto_revogado):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'texto-revogado'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_indice_remissivo):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'indice-remissivo'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_titulo_regulamento):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'titulo-regulamento'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_norma_regulamento):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'norma-regulamento'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_nota_de_consolidacao):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'nota-de-consolidacao'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph if validate_data(text_paragraph, rules_and_definitions.pattern_observacao):
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'observacao'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

        case text_paragraph:
            paragraph_classified['Index'] = pos_no_doc
            paragraph_classified['Class'] = 'nao-classificado'
            paragraph_classified['Content'] = text_paragraph
            paragraphs_RICMS.append(paragraph_classified)
            return

def table_to_html(table):
    # Converter a tabela em um DataFrame
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    df = pd.DataFrame(data)

    # Converter DataFrame para HTML
    html_table = df.to_html(index=False, header=False, escape=False)

    return html_table

def create_summary_list(df):
    classes_for_summary = ['capitulo', 'secao','titulo','subsecao','livro','titulo-tabela','anexo']
    classes_for_subtitle = ['titulo-capitulo-secao-ou-subsecao']
    summary_html = ""
    for index, row in df.iterrows():   
        if row['Class'] in classes_for_summary:
            class_name = row['Class'].lower().replace(' ', '-')
            next_paragraph = df.iloc[index + 1]
            if next_paragraph['Class'] in classes_for_subtitle:
                summary_html += f"<li class='{class_name}'><a href='#{index}'>{row['Content']} - {next_paragraph['Content'].lower()}</a></li>\n"
            else:
                summary_html += f"<li class='{class_name}'><a href='#{index}'>{row['Content']}</a></li>\n"
    logging.info('Sumário criado com sucesso!')
    return summary_html

paragraphs_RICMS = list()
paragraph_classified = dict()

# abrir arquivo do RICMS
doc = Document('data/RICMS-Corrido.docx')

# classificar os parágrafos
for index, elemento in enumerate(doc.element.body):
    if elemento.tag.endswith('p'):  # verifica se o elemento é um parágrafo
        paragraph_classified = dict()  # cria um novo dicionário para guardar a classificação do parágrafo

        content_paragraph = elemento.text.strip()
        # Normalização para remover espaços invisíveis
        content_paragraph = re.sub(r'[\u200b\u200c\u200d\uFEFF]', '', content_paragraph)  # remove espaços invisíveis comuns
        content_paragraph = re.sub(r'\s+', ' ', content_paragraph).strip()  # normaliza espaços em branco
        content_paragraph = content_paragraph.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')

        classify_paragraph(content_paragraph, index)  # classifica o parágrafo a partir das expressões regulares

    elif elemento.tag.endswith('tbl'):
        paragraph_classified = dict()   # cria um novo dicionário para guardar a classificação do parágrafo

        table_index = sum(1 for e in doc.element.body[:index] if e.tag.endswith('tbl'))  # identifica o índice dessa tabela na lista doc.tables
        table_obj = doc.tables[table_index]  # retorna a tabela correta

        classify_paragraph(table_to_html(table_obj), index)

# criar DataFrame e salvar em Excel
df_paragraphs_RICMS = pd.DataFrame(paragraphs_RICMS)
df_paragraphs_RICMS.to_excel('data/paragrafos_classificados.xlsx')

logging.info('Parágrafos classificados com sucesso!')

# criando um arquivo HTML a partir do DF
# início do código HTML
html_content = f"""<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>RICMS/MT</title>
    <link rel="stylesheet" href="style.css">
    <link rel="shortcut icon" href="img/favicon.ico" type="image/x-icon">
</head>
<body>
    <header class="cabecalho">
        <h1>Secretaria de Estado de Fazenda de Mato Grosso</h1>
        <nav>
            <ul class="menu">
                <li><a href="https://www5.sefaz.mt.gov.br/inicio" target="_blank">Site Oficial</a></li>
                <li><a href="https://www.sefaz.mt.gov.br/spl/portalpaginalegislacao" target="_blank">Portal da Legislação</a></li>
                <li><a href="https://portal.mt.gov.br/" target="_blank">Portal Gov MT</a></li>
            </ul>
        </nav>
        <div class="barra-degrade"></div>
    </header>
    <div id="container-principal">
      <button id="toggle-sidebar">☰</button>
      <nav id="sidebar", class="sidebar-hidden">
      <input type="text" id="search-bar" placeholder="Pesquisar...">
        <ol>
          {create_summary_list(df_paragraphs_RICMS)}
        </ol>
      </nav>
        <main>
"""

# Colocando o texto dentro das TAGs adequadas
for index, row in df_paragraphs_RICMS.iterrows():
    if row['Class'] in rules_and_definitions.html_tags_for_class_paragraph.keys():  # os parágrafos que não estiverem classificações
        # listadas em html_tags_for_class_paragraph não serão adicionados ao HTML
        class_name = row['Class'].lower().replace(' ', '-')  # define o nome da classe que será atribuída ao elemento
        # com base no conteúdo da coluna 'Class'
        tag = rules_and_definitions.html_tags_for_class_paragraph.get(row['Class'])  # procura no dicionário html_tags_for_class_paragraph o
        # valor correspondente a 'Key' que é a classe do parágrafo

        html_content += f"<{tag} class='{class_name}' id='{index}'>{row['Content']}</{tag}>\n"

html_content += """
        </main>
    </div>
<footer>
    <p>Site desenvolvido pela UNERC SEFAZ/MT - <a href="https://www.portaldoconhecimento.mt.gov.br/" target="_blank">Portal do Conhecimento</a></p>
</footer>
    <script src="script.js"></script>
</body>
</html>"""

logging.info('Html criado com sucesso!')

# adicionado TAG STRONG nos artigos
html_content = re.sub(rules_and_definitions.pattern_artigo_com_numero, r'<strong>\g<0></strong>', html_content) # '\g<0>' significa - a
# parte da string que correspondeu à expressão regular informada

logging.info('TAGs adicionadas com sucesso!')

# Salvando em um arquivo HTML
with open("../regulamento-icms-sefazmt/index.html", "w", encoding="utf-8") as file:
    file.write(html_content)

logging.info("Arquivo HTML gerado com sucesso!")